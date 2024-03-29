from calendar import c
import os
import string
import time
import re
import json
import yaml
from datetime import datetime

import selenium
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support import expected_conditions
from bs4 import BeautifulSoup

import urllib.parse
# from urllib.parse import urlparse
import logging

from collections import defaultdict

from docx import Document

from openpyxl import load_workbook
from openpyxl.utils import column_index_from_string, get_column_letter

from PIL import Image
from PIL import Image

from Tool.Tool_Web import *
from Tool.Tool_Data import *
from Tool.Tool_SQL import *


class AmazonInfo():
    def __init__(
        self, driver: webdriver.Chrome, wait: WebDriverWait, actions: ActionChains
    ):
        self.driver, self.wait, self.actions = driver, wait, actions
        #self.static_value()
        
        self.projectroot = os.path.dirname(os.path.abspath(__file__))
        parent_directory = os.path.dirname(self.projectroot)
        
        # 生成当前时间的字符串，格式为 YYYYMMDD_HHMM
        current_time = datetime.now().strftime("%Y%m%d_%H%M")
        # 设置文档文件名为当前时间
        self.docfilename = f"Garb_Info_output_{current_time}.docx"
        self.output_root = os.path.join(parent_directory, '# OUTPUT', os.path.basename(self.projectroot))
        self.docfilepath = os.path.join(self.output_root, self.docfilename)
        
        self.listing_root = os.path.join(parent_directory, '# LISTING')
        self.product_root = os.path.join(self.listing_root, '产品数据')
        
        self.e汇总_path = os.path.join(self.listing_root, '汇总表格', 'ASIN_信息汇总.xlsx')
        self.e汇总_wb = load_workbook(filename=self.e汇总_path, read_only=False)
        self.e汇总_Sheet1 = self.e汇总_wb['Sheet1']
        self.e汇总_colstr_链接 = self.find_colname_letter(sheet=self.e汇总_Sheet1, rowindex=2, colname='链接')
        
        self.e队列_path = os.path.join(self.listing_root, '汇总表格', 'ASIN_抓取队列.xlsx')
        self.e队列_wb = load_workbook(filename=self.e队列_path, read_only=False)
        self.e队列_Sheet1 = self.e队列_wb['Sheet1']
        self.e队列_colstr_链接 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='链接')
        self.e队列_colstr_ASIN = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='ASIN')
        self.e队列_colstr_国家 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='国家')
        self.e队列_colstr_有效 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='有效')
        self.e队列_colstr_立即更新 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='立即更新')
        self.e队列_colstr_是否更新 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='是否更新')
        self.e队列_colstr_更新时间 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='更新时间')
        self.e队列_colstr_更新周期 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='更新周期')
        self.e队列_colstr_主图450 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='主图450')
        self.e队列_colstr_主图1500 = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='主图1500')
        self.e队列_colstr_isKeepa = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='isKeepa')
        self.e队列_colstr_isSeller = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='isSeller')
        self.e队列_colstr_isXiYou = self.find_colname_letter(sheet=self.e队列_Sheet1, rowindex=1, colname='isXiYou')
        
        self.x记录_path = os.path.join(self.listing_root, '汇总表格', 'EXCE_表格记录.xlsx')
        self.x记录_wb = load_workbook(filename=self.x记录_path, read_only=False)
        self.x记录_Sheet1 = self.x记录_wb['Sheet1']
        self.x记录_colstr_表格名称 = self.find_colname_letter(sheet=self.x记录_Sheet1, rowindex=1, colname='表格名称')
        self.x记录_colstr_分类 = self.find_colname_letter(sheet=self.x记录_Sheet1, rowindex=1, colname='分类')
        self.x记录_colstr_子类 = self.find_colname_letter(sheet=self.x记录_Sheet1, rowindex=1, colname='子类')
        self.x记录_colstr_ASIN计数 = self.find_colname_letter(sheet=self.x记录_Sheet1, rowindex=1, colname='ASIN计数')
        self.x记录_colstr_更新时间 = self.find_colname_letter(sheet=self.x记录_Sheet1, rowindex=1, colname='更新时间')
        
        self.info_all_count = 0
        self.info_fail_count = 0
        self.doc = Document()

    # !doc日志文档写入
    def append_line(self, line, flag=True):
        # 当 savelog=True 且 flag=True 时，整个表达式为 True。
        # 当 savelog=True 且 flag=False 时，整个表达式为 False。
        # 当 savelog=False 且 flag=True 时，整个表达式为 True。
        # 当 savelog=False 且 flag=False 时，整个表达式为 False。
        # if self.savelog and flag or (not self.savelog and flag):
        if flag:
            # 追加一行到文档中
            self.doc.add_paragraph(line)
            # 保存文档
            self.doc.save(self.docfilepath)
        print(line)

    def append_dict(self, dict):        
        # 迭代字典并添加内容
        for key, value in dict.items():
            self.doc.add_paragraph(f"{key}: {value}")
        # 保存文档
        self.doc.save(self.docfilepath)

    # !获取div特征
    def get_div_features(self, parent, level=0,flag=False):
        # 获取父元素的标签名
        tag_name = parent.tag_name
        # 获取父元素的class属性
        attributes = parent.get_attribute("class")
        # 获取父元素的父元素
        pre_parent = parent.find_element(By.XPATH, '..')
        # 获取父元素的所有同类型的同级元素
        elements = pre_parent.find_elements(By.TAG_NAME, parent.tag_name)
        i = 0
        # 递归时，当前步骤中的parent不是第一次输入的值时
        if level != 0:
            # 遍历同级元素列表
            for i, elem in enumerate(elements):
                # 当遍历到当前元素时，记录当前元素在同级元素中的位置
                if elem != parent:
                    i = i + 1
                else:
                    break
        # 标签类型，层级，该层级下有多少个相关元素，class的值，[]下一级元素内容
        feature = [tag_name, level, i, attributes, []]
        if flag:
            space_interval = level*"  "
            self.append_line(f'{space_interval}{feature}')
        # 遍历当前div元素的子元素
        for child in parent.find_elements(By.XPATH, './*'):
            # 对每个子元素递归调用get_div_features()函数，获取子元素的特征
            child_feature = self.get_div_features(child, level + 1, flag)
            # 将子元素的特征添加到当前div元素的特征列表中
            feature[4].append(child_feature)
        return feature

    # !比对div的特征值，找到配置文件中对应的section，在section中根据数据位置获取数据
    # 如果特征值在配置文件中没有找到，则保存特征值
    def match_feature_data(self, div_Info_child, config_name):
        config_file = f'yaml/features_result_{config_name}.yml'
        # 获取存放图片目录路径
        config_img = os.path.join(self.output_root, 'yaml', f'img_{config_name}')
        
        self.append_line(f'&&当前ASIN下的div元素有：{len(div_Info_child)}个')
        with open(config_file) as f:
            config = yaml.safe_load(f)
            
        # 遍历div_Info_child数组
        for index, child in enumerate(div_Info_child, 1):
            # 当文本为空或空字串时继续下一个循环
            if child.text in None or child.text == '': continue
            # 将数字转字符串
            index_str = f'0{index}' if index < 10 else str(index)
            child_class = child.get_attribute("class")
            print(child_class)
            feature_symbol = None
            if child_class == '':
                feature_symbol = 'div'
            child_xpath = self.get_xpath(child)
            self.append_line(f'--当前匹配的是第 {index_str} 个div')
            self.append_line(f'--当前匹配的xpath：{child_xpath}')
            # 先获取div的特征值
            features_list = self.get_div_features(child, 0)
            # 将特征值转为json
            feature_str = json.dumps(features_list)
            exists = False
            self.info_all_count += 1
            key_value = None
            value = None
            # 遍历配置文件
            for section in config:
                # 匹配特征值
                if config[section]['Div_feature'] == feature_str:
                    self.count_section(section)
                    self.append_line(f'--已匹配对应的特征值section：{section}')
                    exists = True
                    key_flag = 1
                    while key_flag > 0:
                        # section中有几个需要获取的data，后缀就到几
                        key_name = f'data_{key_flag}'
                        # 匹配失败，即section的需要获取值已经到尾部，跳出循环
                        if key_name not in config[section]:
                            key_flag = -1
                            continue
                        # 匹配成功后，获取键值
                        key_value = config[section][key_name]
                        data_name = key_value[0]
                        data_method = key_value[1]
                        data_xpath = key_value[2]
                        data_type = key_value[3]
                        target_div = child.find_element(By.XPATH, data_xpath)
                        # 获取值的方式
                        # 如果data_type=hiddentext，通过js代码修改元素的class，再获取值
                        if data_method == 'attribute':
                            value = target_div.get_attribute(data_type)
                        elif data_method == 'xpath':
                            if data_type == 'hiddentext':
                                self.driver.execute_script(
                                    "arguments[0].className = '';", target_div
                                )
                            value = target_div.text
                        self.append_line(f'匹配到的：{section}\t_数据名称：{data_name}\t_数据值=>{value}')
                        # 通过data_name找到self中对应名称的值，将value值赋值给self.'data_name'
                        # 如果 self 之前没有叫做 data_name 的属性，这条语句将会给 self 增加一个新的属性，并将 value 赋值给它
                        setattr(self, data_name, value)
                        key_flag += 1
                    break  # exists = True
            new_section = None
            if not exists:
                self.get_element_structure(child)
                # 重新获取一次div特征值，并打印在文档中
                self.get_div_features(child, 0, True)
                self.info_fail_count += 1
                # 'Section{}'.format(len(config)+1) len(config)+1 就是数量+1,即新增一个section后的总数，'Section{}'.format(3) = 'Section3'
                #new_section = f'{feature_symbol} data_i.{info.data_index} c.{index}'
                new_section = f'data_i.{self.ASIN}_c.{index_str}'                # 当 data_index 为 None时，该div时HR下的asin
                config[new_section] = {'Div_feature': feature_str}
                self.append_line(f'!!新增元素特征值：{new_section}\t目标配置类型：{config_name}')
                with open(config_file, 'w') as f:
                    yaml.dump(config, f)
                    
                y = child.location['y']
                self.driver.execute_script(f"window.scrollTo(0, {y});")
                # 获取元素的x,y坐标
                child_x = child.location['x']
                child_y = child.location['y']
                child_width = child.size['width']
                child_height = child.size['height']
                # 截取指定区域
                self.driver.get_screenshot_as_file('screenshot.png')
                img = Image.open('screenshot.png')
                #img.save(f'{config_img}\\screenshot.png')                
                img_crop = img.crop((child_x, 0, child_x+child_width, child_height))
                try:
                    img_crop.save(f'{config_img}\\{new_section}.png')
                except Exception:
                    self.append_line(f'{new_section}\t @@child_width:{child_width}=child_height:{child_height}\t 目标为空div')

    # !计算section出现的次数
    def count_section(self, section_name):
        '''
        if section_name in self.section_dict:
            self.section_dict[section_name] += 1
        else:
            self.section_dict[section_name] = 1
        '''
        self.section_dict = defaultdict(int)
        self.section_dict[section_name] += 1

    # !格式化html并输出到doc
    def get_element_structure(self, element, level=0):
        outer_html = element.get_attribute('outerHTML')
        soup = BeautifulSoup(outer_html, "html.parser")
        # 使用prettify()函数将HTML代码美化
        pretty_html = soup.prettify()
        # 打印美化后的HTML代码，每行打印一次
        for line in pretty_html.split('\n'):
            self.append_line(line)

    # !获取div的xpath
    def get_xpath(self, element):
    # 将JavaScript getXpath函数嵌入到Python代码中
        script = """
        function getElementXPath(element) {
            if (element && element.nodeType === Node.ELEMENT_NODE) {
                var xpath = '';
                var parent = null;
                for (parent = element; parent && parent !== document; parent = parent.parentNode) {
                    var index = 1;
                    for (previousSibling = parent.previousSibling; previousSibling; previousSibling = previousSibling.previousSibling) {
                        if (previousSibling.nodeName === parent.nodeName) {
                            index++;
                        }
                    }
                    xpath = '/' + parent.nodeName + '[' + index + ']' + xpath;
                }
                return xpath;
            } else {
                throw new Error('Invalid input: Element is not a valid DOM node');
            }
        }
        return getElementXPath(arguments[0]);
        """
        return self.driver.execute_script(script, element)

    # !找到列名对应的列序号，返回字母
    def find_colname_letter(self, sheet, rowindex, colname, match_mode='精准匹配'):
        # next：这个函数会返回一个迭代器的下一个元素。
        # next 用于获取满足条件（该行的值等于colname）的第一个元素的列字母。如果没有元素满足条件，它将返回一个默认值，这里是None
        #cell_letter = None
        #for cell in sheet[rowindex]:
        #    if  match_mode == '精准匹配':
        #        if cell.value == colname:
        #            cell_letter = cell.column_letter
        #    elif match_mode == '模糊匹配':
        #        if colname in cell.value:
        #            cell_letter = cell.column_letter
        # 等价于：
        cell_letter = next(
            (
                cell.column_letter
                for cell in sheet[rowindex]
                if  (match_mode == '精准匹配' and cell.value == colname) or (match_mode == '模糊匹配' and colname in cell.value)
            ),
            None,)
        return cell_letter

    # !格式化链接
    # https://www.amazon.com/sspa/click?ie=UTF8&spc=MTo2Mjk2MDY0NTc3MDY0MDg3OjE3MTE2MjQyMTI6c3BfYXRmOjMwMDAzMTY0Nzg2NzQwMjo6MDo6&url=%2FDinosaur-Storage-Educational-Realistic-Jurassic%2Fdp%2FB0B3RRZX8R%2Fref%3Dsr_1_1_sspa%
    # 将以上的链接转缓存https://www.amazon.com/dp/{asin} 其中{asin}是从链接中找到的asin，上面链接中的{asin}是B0B3RRZX8R
    # dp%2F后面的十个字符是asin，
    # 输出 return https://www.amazon.com/{asin}, {asin}, 国家 
    # 当url的域名后缀是com，国家是us
    # 当url的域名后缀是co.uk，国家是uk
    # 当url的域名后缀是其他是，国家就是后缀
    def format_url(self, url):
        if 'sspa' in url:
            asin_pattern = re.compile(r'dp%2F([A-Za-z0-9]{10})')
        else:
            asin_pattern = re.compile(r'/dp/([A-Z0-9]{10})')
        match = asin_pattern.search(url)
        asin = match[1] if match else None
        parsed_url = urllib.parse.urlparse(url)
        domain = parsed_url.netloc
        domain_parts = domain.split('.')
        if len(domain_parts) >= 4:
            domain_suffix = f'{domain_parts[-2]}.{domain_parts[-1]}'
        else:
            domain_suffix = domain_parts[-1]
        domain_suffix_country_dict = {'com': 'us', 'co.uk': 'uk'}
        country = domain_suffix_country_dict.get(domain_suffix, domain_suffix)
        new_url = f'https://{domain}/dp/{asin}'
        return new_url, asin, country

    # !目标报价表：找到表的某一列中符合该值的行号
    def find_colindex_value_rowindex(self, sheet, rowvalue, colindex, match_mode='精准匹配'):        
        for index, row in enumerate(sheet.iter_rows(), start=1):
        # 如果匹配模式为 '精准匹配'
            if match_mode == '精准匹配':
                if row[colindex-1].value == rowvalue: # row[0] 获取当前列的值
                    return index # 返回匹配行的行号
            # 如果匹配模式为 '模糊匹配'
            elif match_mode == '模糊匹配':
                if rowvalue in str(row[colindex-1].value):
                    return index # 返回匹配行的行号
        # 如果没有找到匹配的行，返回 None
        return None

    # !计算表格的总行数和总列数
    def tool_count(self,sheet):
        row_count = 0
        while sheet.cell(row=row_count+1, column=1).value is not None:
            row_count += 1
        column_count = 0
        while sheet.cell(row=1, column=column_count+1).value is not None:
            column_count += 1       
        return row_count, column_count

    # !遍历目标文件夹下对应后缀的文件
    def list_files_by_type(self, directory , file_type='.xlsx'):
        excel_files = []
        # 对指定目录及其所有子目录进行遍历
        for root, dirs, files in os.walk(directory):
            for file in files:
                if file.endswith(file_type):
                    excel_files.append(os.path.join(root, file))
        return excel_files

    # !整合表格数据
    def excel_merge_with(self):  
        file_list = self.list_files_by_type(self.product_root, '.xlsx')
        for file_path in file_list:
            ase_name = os.path.basename(file_path)
            dir_name = os.path.basename(os.path.dirname(file_path))
            pdir_name = os.path.basename(os.path.dirname(os.path.dirname(file_path)))
            # 如果只有一级目录，值向上传递一次
            if pdir_name == '产品数据':
                pdir_name = dir_name
                dir_name = None
            # any：这个函数测试可迭代的元素是否有至少一个为真。
            # any用于检查表格中是否有至少一个单元格的值等于ase_name。如果有任何一个单元格的值等于 ase_name, 则 any 函数返回 True，否则返回 False。
            wb = load_workbook(filename=file_path,read_only=True)
            ws = wb['Sheet1']
            ws_maxrow, ws_maxcol = self.tool_count(ws)
            x记录_Sheet1_maxrow, x记录_Sheet1_maxcol= self.tool_count(self.x记录_Sheet1)
            ase_name_exist = any(
                cell.value == ase_name
                for cell in self.x记录_Sheet1[self.x记录_colstr_表格名称]
            )
            # 更新ASIN队列信息
            self.excel_write_link(file_path)
            if not ase_name_exist:
                self.x记录_Sheet1[f'{self.x记录_colstr_表格名称}{x记录_Sheet1_maxrow+1}'] = ase_name
                self.x记录_Sheet1[f'{self.x记录_colstr_分类}{x记录_Sheet1_maxrow+1}'] = pdir_name
                self.x记录_Sheet1[f'{self.x记录_colstr_子类}{x记录_Sheet1_maxrow+1}'] = dir_name
                self.x记录_Sheet1[f'{self.x记录_colstr_ASIN计数}{x记录_Sheet1_maxrow+1}'] = ws_maxrow - 1
                self.x记录_Sheet1[f'{self.x记录_colstr_更新时间}{x记录_Sheet1_maxrow+1}'] = datetime.now().strftime("%Y/%m/%d %H:%M")
            # 更新ASIN计数
            # ?当asin需要立即更新时，asin计数是不变的。所以这里只记录新增asin后的更新时间
            else:
                所在行号 = self.find_colindex_value_rowindex(self.x记录_Sheet1, ase_name, column_index_from_string(self.x记录_colstr_表格名称))
                ASIN计数 = self.x记录_Sheet1.cell(row=所在行号, column=column_index_from_string(self.x记录_colstr_ASIN计数)).value
                if ASIN计数 < x记录_Sheet1_maxrow - 1:
                    self.x记录_Sheet1[f'{self.x记录_colstr_ASIN计数}{所在行号}'] = ws_maxrow - 1
                    self.x记录_Sheet1[f'{self.x记录_colstr_更新时间}{所在行号}'] = datetime.now().strftime("%Y/%m/%d %H:%M")
                    
        self.x记录_wb.save(self.x记录_path)

    # !将表格中的链接都汇总到抓取队列中
    # 如果有一个ASIN想重点更新  更新=2 则更新周期为2
    def excel_write_link(self, file_path):
        wb = load_workbook(filename=file_path, read_only=False)
        ws = wb['Sheet1']
        colstr_链接 = self.find_colname_letter(sheet=ws, rowindex=1, colname='链接')
        colstr_ASIN = self.find_colname_letter(sheet=ws, rowindex=1, colname='ASIN')
        colstr_国家 = self.find_colname_letter(sheet=ws, rowindex=1, colname='国家')
        colstr_更新 = self.find_colname_letter(sheet=ws, rowindex=1, colname='更新')
        colnum_链接 = column_index_from_string(colstr_链接) -1
        colnum_ASIN = column_index_from_string(colstr_ASIN) -1
        colnum_国家 = column_index_from_string(colstr_国家) -1
        colnum_更新 = column_index_from_string(colstr_更新) -1
        ws_maxrow, ws_maxcol = self.tool_count(ws)
        # 最低从第二行开始，index的起始值是2
        for index, row in enumerate(ws.iter_rows(min_row=2, max_row=ws_maxrow, max_col=4),start=2):
            链接 = row[colnum_链接].value
            asin = row[colnum_ASIN].value
            国家 = row[colnum_国家].value
            更新 = row[colnum_更新].value
            链接_exist = any(
                cell.value == 链接
                for cell in self.e队列_Sheet1[self.e队列_colstr_链接]
            )
            # 为空时立即添加到ASIN队列
            if not 链接_exist:
                链接,asin,国家 = self.format_url(链接)
                ws[f'{colstr_链接}{index}'] = 链接
                ws[f'{colstr_ASIN}{index}'] = asin
                ws[f'{colstr_国家}{index}'] = 国家
                ws[f'{colstr_更新}{index}'] = 7
                e队列_Sheet1_maxrow, e队列_Sheet1_maxcol = self.tool_count(self.e队列_Sheet1)
                self.e队列_Sheet1[f'{self.e队列_colstr_链接}{e队列_Sheet1_maxrow+1}'] = 链接
                self.e队列_Sheet1[f'{self.e队列_colstr_ASIN}{e队列_Sheet1_maxrow+1}'] = asin
                self.e队列_Sheet1[f'{self.e队列_colstr_国家}{e队列_Sheet1_maxrow+1}'] = 国家
                self.e队列_Sheet1[f'{self.e队列_colstr_立即更新}{e队列_Sheet1_maxrow+1}'] = True
                self.e队列_Sheet1[f'{self.e队列_colstr_是否更新}{e队列_Sheet1_maxrow+1}'] = True
                self.e队列_Sheet1[f'{self.e队列_colstr_更新周期}{e队列_Sheet1_maxrow+1}'] = 7
                self.e队列_Sheet1[f'{self.e队列_colstr_主图450}{e队列_Sheet1_maxrow+1}'] = True
                self.e队列_Sheet1[f'{self.e队列_colstr_主图1500}{e队列_Sheet1_maxrow+1}'] = True
                self.e队列_Sheet1[f'{self.e队列_colstr_isKeepa}{e队列_Sheet1_maxrow+1}'] = True
                self.e队列_Sheet1[f'{self.e队列_colstr_isSeller}{e队列_Sheet1_maxrow+1}'] = True
                self.e队列_Sheet1[f'{self.e队列_colstr_isXiYou}{e队列_Sheet1_maxrow+1}'] = True
            # 如果更新==i 立马更新，等表格通过excel_loop_grab填充后，再恢复成e队列_Sheet1中的更新周期
            elif str(更新) == 'i':
                所在行号 = self.find_colindex_value_rowindex(self.e队列_Sheet1, 链接, column_index_from_string(self.e队列_colstr_链接))
                self.e队列_Sheet1[f'{self.e队列_colstr_立即更新}{所在行号}'] = True
            # 更新ASIN队列的更新周期
            else:
                所在行号 = self.find_colindex_value_rowindex(self.e队列_Sheet1, 链接, column_index_from_string(self.e队列_colstr_链接))
                self.e队列_Sheet1[f'{self.e队列_colstr_更新周期}{所在行号}'] = 更新
                
        wb.save(file_path)
        self.e队列_wb.save(self.e队列_path)

    def excel_loop_grab(self):
        # 最低从第三行开始，index的起始值是3
        maxrow, maxcol= self.tool_count(self.e队列_Sheet1)
        for index, row in enumerate(self.e队列_Sheet1.iter_rows(min_row=3,max_row=maxrow,max_col=maxcol), start=3):
            
            pass

    def excel_fill_in(self):
        pass

    # !
    def garb_info(self, url):  
        #TODO:self的值
        info_主要信息 = self.driver.find_element(By.XPATH, '//*[@id="centerCol"]')
        info_主要信息_child = info_主要信息.find_elements(By.XPATH, './div')
        self.match_feature_data(info_主要信息_child, 'asin')
        
        #?限购 如何判断
        
        #Product information
        #//*[@id="productDetails_detailBullets_sections1"]
        #键值的空格转换成下划线_
        
        #TODO:关联ASIN
        #class="a-row"
        #先获取h2且class=a-carousel-heading a-inline-block
        #//*[contains(@value, 'a-carousel-heading a-inline-block')]
        #sp_desktop_sponsored_label
        #?如果当该div存在时，使用赋值 例如 method=value value=true
        # <span class="a-carousel-page-current">1</span>
        # <span class="a-carousel-page-max">4</span>
        
        # 向上遍历到
        # id="similarities_feature_div" class="celwidget" data-feature-name="similarities" data-csa-c-type="widget" data-csa-c-content-id="similarities" data-csa-c-slot-id="similarities_feature_div"
        # id="similarities_feature_div" class="celwidget" data-feature-name="similarities" data-csa-c-type="widget" data-csa-c-content-id="similarities" data-csa-c-slot-id="similarities_feature_div"
        # 停止遍历到 且不执行找ol的步骤
        # id="rhf" class="copilot-secure-display" style="clear: both;" role="complementary" aria-label="Your recently viewed items and featured recommendations">
        # <div class="rhf-frame" style="display: block;">
        
        # <div class="a-carousel-col a-carousel-left"
        # <a class="a-button a-button-image a-carousel-button a-carousel-goto-prevpage" tabindex="0" href="#" id="a-autoid-47"
        
        # <div class="a-carousel-col a-carousel-center">
        # <div class="a-carousel-viewport" id="anonCarousel1"
        # <ol class="a-carousel" role="list" style="width: 1307px;">
        # 到最后一页时，再次点击会返回第一页
        
        # <div class="a-carousel-col a-carousel-right" style="height: 308px; visibility: visible;">
        # <a class="a-button a-button-image a-carousel-button a-carousel-goto-nextpage" tabindex="0" href="#" id="a-autoid-48" aria-disabled="false" 
        try:
            info_经常购买 = self.driver.find_element(By.XPATH, '//*[@id="CardInstanceAIsiVJYo9vE1IudOdPrv2Q"]/div/div[1]')
            info_四星产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel1"]/ol')
            # Products related to this item 在 What's in the box 和 Videos中间/4 star 上面
            info_四星产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel2"]/ol')
            # From the brand 品牌故事的展示信息
            info_四星产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel3"]/ol')
            # Videos 相关产品的视频
            info_四星产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel4"]/ol')
            # Similar brands on Amazon
            info_四星产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel5"]/ol')
            # Reviews with images
            info_四星产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel6"]/ol')
            # Products related to this item 在 Product information 上面
            info_相关产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel7"]/ol')
            # Related Climate Pledge Friendly items
            info_相关产品 = self.driver.find_element(By.XPATH, '//*[@id="anonCarousel8"]/ol')
            # Related products with free delivery on eligible orders
        except Exception:
            print (Exception)
        
        #TODO:卖家精灵
        # 卖家精灵：抢登录/选择登录账号-Exception
        # 卖家精灵：采集排名和排名截图
        # 卖家精灵：AI评论分析
        # 卖家精灵：Q&A下载


    #update方法可以将返回的字典合并到result字典中
    def get_title(self):
        # 标题、价格 获取//*[@id="expandTitleToggle"]、//*[@id="corePriceDisplay_desktop_feature_div"]/div[1]/span[1]/span[1] 中的文本
        self.title = self.driver.find_element(By.XPATH, '//*[@id="productTitle"]').text
        return {'标题':self.title}

    def get_price(self):
        # sourcery skip: extract-duplicate-method, hoist-statement-from-if
        sale_price = None
        rrp_price = None
        is_Deal = False
        discount = None
        
        # 检查是否可售
        isAvailable = False
        try:
            unavailable_element = self.driver.find_element(
                By.XPATH, '//*[@id="availability"]/span'
            )
            # 转化成小写
            unavailable_text = unavailable_element.text.lower()
            print(unavailable_text)
        except Exception:
            isAvailable = True
        else:
            # 当其他站点  如de isAvailable = True
            if 'unavailable' in unavailable_text:
                sale_price = 'Unavailable'
                isAvailable = False
            elif 'stock' in unavailable_text:
                isAvailable = True
        
        try:
            # 检查是否在Deal //*[@id="dealBadge_feature_div"]/span/span/span
            temp = self.driver.find_element(
                By.XPATH, '//*[@id="dealBadge_feature_div"]/span/span/span'
            ).text
            is_Deal = temp == 'Deal'
        except Exception:
            is_Deal = False

        prime_price = None
        try:
            prime_element = self.driver.find_element(
                By.XPATH, '//*[@id="pep-signup-link"]/span[2]'
            )
            prime_price = prime_element.text
        except Exception:
            print('No prime price element found')
        else:
            print(f'prime price : {prime_price}')

        if isAvailable:
            try:
                # 检查是否有折扣标识
                # apex_desktop 下两级 有的是 apex_desktop_newAccordionRow 有的是 apex_desktop_qualifiedBuybox
                apex_element = self.driver.find_element(By.XPATH, '//*[@id="apex_desktop"]')
                corePrice_element = apex_element.find_element(
                    By.XPATH, '//*[@id="corePriceDisplay_desktop_feature_div"]'
                )
                discount_element = corePrice_element.find_elements(
                    By.XPATH,
                    '//*[@id="corePriceDisplay_desktop_feature_div"]/div[1]/span[contains(@class, "a-size-large a-color-price savingPriceOverride aok-align-center reinventPriceSavingsPercentageMargin savingsPercentage")]',
                )[0]
                discount = discount_element.text
            except Exception:
                print('No discount info')
                # 无折扣,售价在第一个span
                apex_element = self.driver.find_element(By.XPATH, '//*[@id="apex_desktop"]')
                sale_price_element = apex_element.find_element(
                    By.XPATH,
                    '//*[@id="corePriceDisplay_desktop_feature_div"]/div[1]/span[1]/span[1]',
                )
                self.driver.execute_script("arguments[0].className = ''; ", sale_price_element)
                sale_price = sale_price_element.text
            else:
                if '%' in discount_element.text:
                    # 有折扣,售价在第二个span,原价在第四个span
                    sale_price_element = discount_element.find_element(
                        By.XPATH,
                        '//*[@id="corePriceDisplay_desktop_feature_div"]/div[1]/span[2]/span[1]',
                    )
                    # 删除class="a-offscreen"
                    self.driver.execute_script(
                        "arguments[0].className = ''; ", sale_price_element
                    )
                    # 再次获取text,此时可以获得正确值
                    sale_price = sale_price_element.text

                    rrp_price_element = discount_element.find_element(
                        By.XPATH,
                        '//*[@id="corePriceDisplay_desktop_feature_div"]/div[2]/span/span[1]/span/span[1]',
                    )
                    self.driver.execute_script(
                        "arguments[0].className = ''; ", rrp_price_element
                    )
                    rrp_price = rrp_price_element.text
                else:
                    print('No discount info')
                    # 无折扣,售价在第一个span
                    apex_element = self.driver.find_element(By.XPATH, '//*[@id="apex_desktop"]')
                    sale_price_element = apex_element.find_element(
                        By.XPATH,
                        '//*[@id="corePriceDisplay_desktop_feature_div"]/div[1]/span[1]/span[1]',
                    )
                    self.driver.execute_script(
                        "arguments[0].className = ''; ", sale_price_element
                    )
                    sale_price = sale_price_element.text

            print(f'Is In Deal: {is_Deal}')
            print(f'Sale Price: {sale_price}')
            print(f'RRP Price: {rrp_price}') if rrp_price else print('No RRP Price')
            print(f'Discount off: {discount}')
        return {'isDeal': is_Deal,'现价':sale_price,'RRP':rrp_price,'折扣':discount}

    def get_rating(self):
        rating = ''
        review = ''
        try:
            parent = self.driver.find_element(By.XPATH, '//*[@id="averageCustomerReviews"]')
            rating = parent.find_element(
                By.XPATH, '//*[@id="acrPopover"]/span[1]/a/span'
            ).text

            review = parent.find_element(By.XPATH, '//*[@id="acrCustomerReviewText"]').text

        except Exception:
            print('No rating review element found')
        else:
            try:
                review = review.replace('ratings', '').strip()
            except Exception:
                print('')
            print(f'rating : {rating}')
            print(f'review : {review}')
        return {'评分':rating, '评价数':review}

    def get_brand(self): 
        brand = ''
        try:
            brand = self.driver.find_element(By.ID, 'bylineInfo').text
            # 过滤字符串
            brand = brand.replace(':', '')  # 去除冒号
            brand = brand.replace('Visit the', '')  # 去除Visit the
            brand = brand.replace('Store', '')  # 去除Store
            brand = brand.replace('Brand', '')  # 去除Brand:
            brand = brand.strip()  # 去除首尾空格
            print(brand)
        except Exception:
            print('No brand element found')
        else:
            print(f'brand : {brand}')
        return {'品牌': brand}

    def get_variant(self):
        variant_info = []
        try:
            variant_element = self.driver.find_element(By.ID, 'twister_feature_div')
            variant_divs = variant_element.find_elements(
                By.XPATH,
                './/div[contains(@class, "twisterSwatchWrapper_0 twisterSwatchWrapper twisterImages thinWidthOverride")]',
            )
            for variant_div in variant_divs:
                # ./ancestor::li[1] 来相对于 variant_div 元素查找最近的 li 父元素
                li_element = variant_div.find_element(By.XPATH, './ancestor::li[1]')
                variant_ASIN = li_element.get_attribute('data-defaultasin')
                # 添加了一个.，表示相对于variant_div元素进行定位。这样可以确保XPath表达式在当前variant_div元素的上下文中执行，而不是全局执行。
                # variant_name = variant_div.find_element(By.XPATH,'.//div[1]/img').get_attribute('src')
                # variant_img = variant_div.find_element(By.XPATH,'.//div[1]/img').get_attribute('src')
                # variant_price = variant_div.find_element(By.XPATH,'.//div[2]/div/span/p').text
                variant_name = ''
                try:
                    variant_name = variant_div.find_element(
                        By.XPATH, './/div[contains(@class, "twisterTextDiv")]//p[last()]'
                    ).text
                except Exception:
                    print('No variant_name found')
                variant_img = ''
                try:
                    variant_img = variant_div.find_element(
                        By.XPATH, './/div[contains(@class, "twisterImageDiv")]//img[last()]'
                    ).get_attribute('src')
                except Exception:
                    print('No variant_img foQund')
                variant_price = variant_div.find_element(
                    By.XPATH, './/div[contains(@class, "twisterSlotDiv")]//p[last()]'
                ).text
                print(
                    f'variant_ASIN:{variant_ASIN}====variant_img:{variant_img}=====variant_price:{variant_price}'
                )
                variant_info.append(
                    {
                        'ASIN': variant_ASIN,
                        'name': variant_name,
                        'img': variant_img,
                        'price': variant_price,
                    }
                )
        except Exception:
            print('No variant element found')
        else:
            variant_count = len(variant_divs)
            print(f'variant_info : {variant_info}')
            print(f'variant_count : {variant_count}')
        return {'变体':variant_info}

    def get_coupon(self):
        coupon = None
        saving = None
        promotion = None
        # 找到父元素//*[@id="promoPriceBlockMessage_feature_div"]
        parent = self.driver.find_element(
            By.XPATH, '//*[@id="promoPriceBlockMessage_feature_div"]'
        )
        try:
            # 优惠券 找到子元素1//*[@id="couponTextpctch*"](用contain匹配)
            coupon = parent.find_element(
                By.XPATH, '//*[@id[contains(., "couponTextpctch")]]'
            ).text
            # coupon = coupon.replace('|', '')
            # coupon = coupon.replace('Terms', '')
            # coupon = coupon.strip()
            coupon = re.findall(r'[\d%]+', coupon)
        except Exception:
            print('No coupon element found')
        else:
            print(f'coupon : {coupon}')
        
        try:
            # 优惠2a 找到子元素2//*[@id="couponBadgepctch*"](用contain匹配)
            saving_a = parent.find_element(
                By.XPATH, '//*[@id[contains(., "couponBadgepctch")]]'
            ).text
            # 优惠2b 找到子元素3//*[@id="promoMessagepctch*"](用contain匹配)
            saving_b = parent.find_element(
                By.XPATH, '//*[@id[contains(., "promoMessagepctch")]]'
            ).text
            saving = f"{saving_a} {saving_b}"
            saving = saving.replace('|', '')
            saving = saving.replace('Terms', '')
            saving = saving.strip()
        except Exception:
            print('No saving element found')
        else:
            print(f'saving : {saving}')
        
        try:
            # 父元素//*[@id="applicablePromotionList_feature_div"] 与优惠1 2 不同
            # 优惠券3 找到元素//*[@id="applicable_promotion_list_sec"]/span/span/a/span[2]/span[1]
            parent = self.driver.find_element(
                By.XPATH, '//*[@id="applicablePromotionList_feature_div"]'
            )
            promotion_a = parent.find_element(
                By.XPATH,
                '//*[@id="applicable_promotion_list_sec"]/span/span/a/span[2]/span[1]',
            ).text
            promotion_b = parent.find_element(
                By.XPATH,
                '//*[@id="applicable_promotion_list_sec"]/span/span/a/span[2]/span[2]',
            ).text
            promotion = f"{promotion_a} {promotion_b}"
            promotion = promotion.replace('|', '')
            promotion = promotion.replace('Terms', '')
            promotion = promotion.strip()
        except Exception:
            print('No promotion element found')
        else:
            print(f'promotion : {promotion}')

        return {'coupon':coupon,'saving':saving,'promotion':promotion}

    def get_amzchoice(self):
        amz_choice = None
        try:
            amz_choice = self.driver.find_element(
                By.XPATH, '//*[@id="acBadge_feature_div"]/div/span[2]/span/span/a'
            ).text
        except Exception:
            print('No Amazon Choice element found')
        else:
            print(f'Amazon Choice : {amz_choice}')
        
        return {'Amazon Choice':amz_choice}

    def get_bullet(self):
        # 五点描述 找到父元素//*[@id="feature-bullets"],获取父元素下//*[@id="feature-bullets"]/ul/li 的li中的文本,保存为数组
        bullet_points = []
        try:
            parent = self.driver.find_element(By.XPATH, '//*[@id="feature-bullets"]/ul')
        except Exception:
            print('No bullet_points element found')
        else:
            lis = parent.find_elements(By.XPATH, './/li')
            bullet_points.extend(li.text for li in lis)
            print(f'bullet_points : {bullet_points}')

        return {'五点描述':bullet_points}

    def get_baseinfo(self):
        # 获取table //*[@id="productOverview_feature_div"]中的信息,用字典保存
        base_info = {}
        try:
            parent = self.driver.find_element(
                By.XPATH, '//*[@id="productOverview_feature_div"]/div/table'
            )
            trs = parent.find_elements(By.XPATH, './/tbody/tr')
            for tr in trs:
                key = tr.find_element(By.XPATH, 'td[1]').text
                value = tr.find_element(By.XPATH, 'td[2]').text
                base_info[key] = value
        except Exception:
            print('No Basic Info element found')
        else:
            print(f'Basic Info : {base_info}')

    def get_productinfo(self):
        # 获取table //*[@id="productDetails_techSpec_section_1"]中的信息,用字典保存
        dict_data_Details = {}
        try:
            table1 = self.driver.find_element(
                By.XPATH, '//*[@id="productDetails_techSpec_section_1"]'
            )
            for tr in table1.find_elements(By.TAG_NAME, 'tr'):
                th = tr.find_element(By.TAG_NAME, 'th')
                td = tr.find_element(By.TAG_NAME, 'td')
                key = th.text
                value = td.text
                dict_data_Details[key] = value
        except Exception:
            print('No dict_data_Details element found')
        else:
            print(f'dict_data_Details : {dict_data_Details}')
        
        return {'展示信息':dict_data_Details}

    def get_datainfo(self):
        # 获取table //*[@id="productDetails_detailBullets_sections1"] 中的信息,用字典保存
        dict_data_Info = {}
        try:
            table2 = self.driver.find_element(
                By.XPATH, '//*[@id="productDetails_detailBullets_sections1"]'
            )
            for tr in table2.find_elements(By.TAG_NAME, 'tr'):
                th = tr.find_element(By.TAG_NAME, 'th')
                td = tr.find_element(By.TAG_NAME, 'td')
                key = th.text
                value = td.text
                dict_data_Info[key] = value
        except Exception:
            print('No dict_data_Info element found')
        else:
            print(f'dict_data_Info : {dict_data_Info}')
        
        return {'产品信息':dict_data_Info}

    

    def get_main450(self):
        span = get_imgspan(self.driver)[0]
            # 使用ActionsChains点击input
        if input is not None:
            self.actions.move_to_element(span)
            self.actions.click(span)
        self.actions.perform()
        time.sleep(0.2)
        # 获取所有450尺寸的主图链接
        image_main450 = []
        ul = self.driver.find_element(By.XPATH, '//*[@id="main-image-container"]/ul')
        for li in ul.find_elements(By.XPATH, 'li'):
            if 'image' in li.get_attribute('class') and 'item' in li.get_attribute( # type: ignore
                'class'
            ): # type: ignore
                img = li.find_element(By.XPATH, './span/span/div/img')
                img_src = img.get_attribute('src')
                image_main450.append(img_src)
        print(image_main450)

        # 遍历父元素下所有元素
        elements = []
        def get_elements(parent):
            children = parent.find_elements(By.XPATH, './*')
            for child in children:
                elements.append(child)
                get_elements(child)
            return elements
        
        return {'主图450':image_main450}

    def get_main1500(self):
        image_left1 = get_imgspan(self.driver)[1]
        image_main1500 = []
        # 如果参数isBigImg为真，获取1000+的大尺寸主图
        try:
            # 先点击第一张主图(侧边栏)
            # image_left1 = self.driver.find_element(By.XPATH, '//*[@id="a-autoid-6"]')
            self.actions.move_to_element(image_left1)
            self.actions.click(image_left1)
            self.actions.perform()
            # 找到图片弹窗元素,模拟点击 //*[@id="imgTagWrapperId"]无法被点击？
            # 部分ASIN 没有[@id="landingImage"] 有class=landingImage
            main_image_element = self.driver.find_element(
                By.XPATH, '//*[@id="main-image-container"]'
            )
            image_popup = main_image_element.find_element(
                By.XPATH, './/div[@class="imgTagWrapper"]/img'
            )
            # image_popup = self.driver.find_element(By.XPATH, '//*[@id="landingImage"]')
            self.actions.move_to_element(image_popup)
            self.actions.click(image_popup)
            self.actions.perform()

            # 等待a-popover-content元素出现 //*[@id="ivImagesTabHeading"]/a
            self.wait.until(
                EC.presence_of_element_located(
                    (By.XPATH, '//*[@id="ivImagesTabHeading"]/a')
                )
            )

            # 找到父元素,和子元素div class="ivRow"
            parent = self.driver.find_element(By.XPATH, '//*[@id="ivThumbs"]')
            divs = parent.find_elements(By.XPATH, './/div[@class="ivRow"]')
            elements = []
            for div in divs:
                des_div = div.find_elements(By.XPATH, './div')
                elements.extend(iter(des_div))
            # 依次点击elements中的元素,并获取图片链接
            for element in elements:
                self.actions.move_to_element(element)
                self.actions.click(element)
                self.actions.perform()
                time.sleep(0.5)
                self.wait.until(
                    EC.presence_of_element_located(
                        (By.XPATH, '//*[@id="ivLargeImage"]/img')
                    )
                )
                img = self.driver.find_element(By.XPATH, '//*[@id="ivLargeImage"]/img')
                img_url = img.get_attribute('src')
                while 'loading' in img_url:  # type: ignore
                    time.sleep(0.5)
                    img_url = img.get_attribute('src')
                image_main1500.append(img_url)
        except Exception:
            print('No description element found')
        else:
            print(f'image_main1500 : ,{image_main1500}')

    def get_video(self):
        # 8.1 获取视频数量
        video_count = 0
        try:
            video_text = self.driver.find_element(By.XPATH, '//*[@id="videoCount"]').text.strip()
            if video_text in ["VIDEO", "VIDEOS"]:
                video_count = 1
            elif 'VIDEOS' in video_text:
                # \d+ 表示匹配一个或多个数字,这是一个正则表达式,而不是一个字符串。所以这里使用转义序列 \d 是正确的,不会产生无效的转义序列错误。
                # 但是,Pylance 分析器误以为这是一个字符串,所以报告了无效的转义序列错误。
                # 在字符串开头添加 r 会让 Pylance 知道这是一个原始字符串,实际上是正则表达式,可以安全地使用转义序列。
                numbers = re.findall(r'\d+', video_text)
                video_count = int(''.join(numbers))
        except Exception:
            print('No videoCount element found')
        else:
            print(f'videoCount : {video_count}')

        return {'视频数量':video_count}

    def get_sellerCapture(self,ASIN:str):
        try:
            self.wait.until(EC.presence_of_all_elements_located((By.XPATH, '//*[@id="quick-view-page"]')))
            seller_parent = self.driver.find_element(By.XPATH,'//*[@id="quick-view-page"]')
            seller_Linechart = self.driver.find_element(By.XPATH,'//*[@id="quick-view-page"]/div[2]/div/div[1]/div[1]/div[2]/span')
            self.actions.move_to_element(seller_Linechart)
            self.actions.click(seller_Linechart)
            self.actions.perform()
            self.wait.until(EC.presence_of_all_elements_located((By.XPATH, '//*[@id="quick-view-page"]/div[2]/div/div[2]/div/div/div/div/div/div[1]/canvas')))
            seller_canvas = seller_parent.find_element(By.XPATH, '//*[@id="quick-view-page"]/div[2]/div/div[2]/div/div/div/div/div/div[1]/canvas')
            seller_selector = seller_parent.find_element(By.CLASS_NAME,'rang-div')
            seller_selector_p = seller_selector.find_elements(By.TAG_NAME,'p')
            if seller_selector_p[-2].find_element(By.XPATH,'./span').text == '最近一年':
                self.actions.move_to_element(seller_selector_p[-2])
                self.actions.click(seller_selector_p[-2])
            else:
                self.actions.move_to_element(seller_selector_p[-1])
                self.actions.click(seller_selector_p[-1])
            self.actions.perform()
            #folder = r'D:\AutoRPA\卖家精灵\ASIN'
            #asin_folder = folder.replace('ASIN', str(ASIN))
            folder = f'D:\\AutoRPA\\卖家精灵\\{ASIN}'
            if not os.path.exists(folder):
                os.makedirs(folder)
            seller_canvas.screenshot(f'{folder}\\seller_canvas.png')
            seller_quickiew = self.driver.find_element(By.XPATH,'//*[@id="seller-sprite-extension-quick-view-listing"]')
            # 获取元素的x,y坐标
            x = seller_quickiew.location['x']
            y = seller_quickiew.location['y']
            # 滚动页面,使元素顶部与页面顶部对齐
            self.driver.execute_script(f"window.scrollTo(0, {y});")
            # 截取指定区域
            self.driver.get_screenshot_as_file('screenshot.png')
            img = Image.open('screenshot.png')
            # 将元素顶部与页面顶部对齐后，y=0
            img = img.crop((x, 0, x+550, 290))
            img.save(f'{folder}\\seller_quickview.png')
        except Exception as e:
            print('No Seller element found')

    def get_sellerRank(self):
        try:
            seller_parent = self.driver.find_element(By.XPATH,'//*[@id="quick-view-page"]')
            # 回滚至顶部
            self.driver.execute_script("window.scrollTo(0, 0);")
            seller_logo = self.driver.find_element(By.XPATH,'//*[@id="quick-view-page"]/div[1]/div[1]/a/img')
            # 悬停至元素上
            self.actions.move_to_element(seller_logo)
            self.actions.perform()
            location = seller_parent.location
            size = seller_parent.size
            print(location)
            print(size)
            self.actions = ActionChains(self.driver)
            # 相对与当前位置偏移
            self.actions.move_by_offset(1412, 180)
            self.actions.click()
            self.actions.perform()
            time.sleep(2)
            # 如何处理下载失败？
        except Exception as e:
            print('No Rank Download Fialure element found')

    def get_description(self):
        # 获取详情页的文本和图片链接
        # 20230720:先进行读取所有文本，再进行除重，最后合成文本
        description = []
        image_description = []
        img_src = ''
        description_srt = ''
        try:
            parent = self.driver.find_element(By.XPATH, '//*[@id="aplus_feature_div"]')
            elements = self.driver.get_elements(parent)
            for element in elements:
                if (
                    element.tag_name == 'p'
                    or element.tag_name.startswith('h')
                    or element.tag_name.startswith('h1')
                    or element.tag_name.startswith('h2')
                    or element.tag_name.startswith('h3')
                    or element.tag_name.startswith('h4')
                    or element.tag_name.startswith('h5')
                    or element.tag_name == 'span'
                ):
                    description.append(element.text)
                elif element.tag_name == 'img':
                    img_src = element.get_attribute('data-src')
                if img_src:
                    image_description.append(img_src)

        except Exception:
            print('No description element found')
        else:
            # 去除重复的元素
            image_description = remove_duplicates(image_description)

            description_srt = ''
            for item in description:
                if item == '':
                    continue
                description_srt += item + '\n'
            description_srt = description_srt.strip()
            print(f'description_srt : {description_srt}')
            print(f'image_description : ,{image_description}')
        
        return {'详情描述':description_srt,'详情图片':image_description}

def remove_duplicates(nums):
    # 使用集合的特性,将数组转换为集合,重复元素会被自动去除
    nums_set = set(nums)
    # 将集合转换回列表
    return list(nums_set)

def get_imgspan(self):
        # 图片 父元素//*[@id="altImages"]/ul
        # 循环点击左边的小图
        ul = self.driver.find_element(By.XPATH, '//*[@id="altImages"]/ul')
        image_left1 = None
        image_flag = 0
        for li in ul.find_elements(By.TAG_NAME, 'li'):
            # 如果li的class包含template或aok-hidden或videoThumbnail,继续下一个循环
            if (
                'template' in li.get_attribute('class')  # type: ignore
                or 'aok-hidden' in li.get_attribute('class')  # type: ignore
                or 'videoThumbnail' in li.get_attribute('class')  # type: ignore
                or 'sellersprite' in li.get_attribute('id')  # type: ignore
            ):
                continue
            span = li.find_element(By.XPATH, './span/span')

            if image_flag == 0:
                image_left1 = span
            image_flag += 1
        
        return [span,image_left1]

# 测试代码
if __name__ == '__main__':
    sc = ChromeStart("Seller")
    driver,wait,actions = sc.GetDriver()
    AmazonI = AmazonInfo(driver,wait,actions)
    #sc.BindPage('https://www.amazon.com/dp/B07H9GY33H',"Contain")
    #AmazonI.get_title()
    AmazonI.garb_info('https://www.amazon.com/dp/B07H9GY33H')
    #print(AmazonI.title) 
    #print(AmazonI.get_price())
    
    #AmazonI.excel_merge_with()
    #AmazonI.excel_write_link('D:\Code\# LISTING\产品数据\玩具-积木\乐高花束\ASIN_Info-乐高花束2.xlsx')
    
