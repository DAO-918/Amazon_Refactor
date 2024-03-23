import os
import time
import re
import json
import yaml
from datetime import datetime
from numbers import Integral

import logging

import selenium
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support import expected_conditions

from bs4 import BeautifulSoup
from urllib.parse import urlparse

from collections import defaultdict

from docx import Document

from PIL import Image
from PIL import Image

import pandas as pd
import openpyxl
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Font

from Tool.Tool_Web import *
from Tool.Tool_Data import *
from Tool.Tool_SQL import *


## ! 可记录搜索结果，asin的位置，即反查排名位置

class Brick4Info():
    def __init__(
        self, driver: webdriver.Chrome, wait: WebDriverWait, actions: ActionChains
    ):
        self.driver, self.wait, self.actions = driver, wait, actions
        self.static_value()
        # 生成当前时间的字符串，格式为 YYYYMMDD_HHMM
        current_time = datetime.now().strftime("%Y%m%d_%H%M")
        # 设置文档文件名为当前时间
        self.docfilename = f"Brick4_Info_Output_{current_time}.docx"
        self.projectroot = os.path.dirname(os.path.abspath(__file__))
        parent_directory = os.path.dirname(self.projectroot)
        self.output_root = os.path.join(parent_directory, '# OUTPUT', os.path.basename(self.projectroot))
        self.docfilepath = os.path.join(self.output_root, self.docfilename)
        self.doc = Document()

    def append_line(self, line, flag=True):
        if flag:
            # 追加一行到文档中
            self.doc.add_paragraph(line)
            # 保存文档
            self.doc.save(self.docfilepath)
        print(line)
        
    def append_dict(self, dict):        
        # 迭代字典并添加内容
        for key, value in dict.items():
            self.doc.add_paragraph(f"{key}: {value}")
        # 保存文档
        self.doc.save(self.docfilepath)

    def get_div_features(self, parent, level=0,flag=False):
        # 获取父元素的标签名
        tag_name = parent.tag_name
        # 获取父元素的class属性
        attributes = parent.get_attribute("class")
        # 获取父元素的父元素
        pre_parent = parent.find_element(By.XPATH, '..')
        # 获取父元素的所有同类型的同级元素
        elements = pre_parent.find_elements(By.TAG_NAME, parent.tag_name)
        i = 0
        # 递归时，当前步骤中的parent不是第一次输入的值时
        if level != 0:
            # 遍历同级元素列表
            for i, elem in enumerate(elements):
                # 当遍历到当前元素时，记录当前元素在同级元素中的位置
                if elem != parent:
                    i = i + 1
                else:
                    break
        # 标签类型，层级，该层级下有多少个相关元素，class的值，[]下一级元素内容
        feature = [tag_name, level, i, attributes, []]
        if flag:
            space_interval = level*"  "
            self.append_line(f'{space_interval}{feature}')
        # 遍历当前div元素的子元素
        for child in parent.find_elements(By.XPATH, './*'):
            # 对每个子元素递归调用get_div_features()函数，获取子元素的特征
            child_feature = self.get_div_features(child, level + 1, flag)
            # 将子元素的特征添加到当前div元素的特征列表中
            feature[4].append(child_feature)
        return feature

    def get_element_structure(self, element, level=0):
        outer_html = element.get_attribute('outerHTML')
        soup = BeautifulSoup(outer_html, "html.parser")
        # 使用prettify()函数将HTML代码美化
        pretty_html = soup.prettify()
        # 打印美化后的HTML代码，每行打印一次
        for line in pretty_html.split('\n'):
            self.append_line(line)

    def get_xpath(self, element):
    # 将JavaScript getXpath函数嵌入到Python代码中
        script = """
        function getElementXPath(element) {
            if (element && element.nodeType === Node.ELEMENT_NODE) {
                var xpath = '';
                var parent = null;
                for (parent = element; parent && parent !== document; parent = parent.parentNode) {
                    var index = 1;
                    for (previousSibling = parent.previousSibling; previousSibling; previousSibling = previousSibling.previousSibling) {
                        if (previousSibling.nodeName === parent.nodeName) {
                            index++;
                        }
                    }
                    xpath = '/' + parent.nodeName + '[' + index + ']' + xpath;
                }
                return xpath;
            } else {
                throw new Error('Invalid input: Element is not a valid DOM node');
            }
        }
        return getElementXPath(arguments[0]);
        """
        return self.driver.execute_script(script, element)
    
    # !!!


# 测试代码
if __name__ == '__main__':
    sc = ChromeStart("Seller", 9222)
    # sc.OpenPage("https://www.amazon.com/")
    sc.BindPage("https://www.amazon.com/s?k=", "Contain")
    driver, wait, actions = sc.GetDriver()
    AmazonS = AmazonSearch(driver, wait, actions)
    # AmazonS.garb_search("remote dinosaur")
    AmazonS.garb_search()

