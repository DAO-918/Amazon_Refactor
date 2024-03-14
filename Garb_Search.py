import os
import time
import re
import json
import yaml
from datetime import datetime

from selenium import webdriver
import selenium
from selenium.webdriver.common.by import By
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support import expected_conditions
from bs4 import BeautifulSoup

from urllib.parse import urlparse

from collections import defaultdict

from docx import Document
import logging

from PIL import Image
from PIL import Image

from Tool.Tool_Web import *
from Tool.Tool_Data import *
from Tool.Tool_SQL import *

## ! 可记录搜索结果，asin的位置，即反查排名位置

class AmazonSearch:
    def __init__(
        self, driver: webdriver.Chrome, wait: WebDriverWait, actions: ActionChains
    ):
        self.driver, self.wait, self.actions = driver, wait, actions
        self.static_value()
        # 生成当前时间的字符串，格式为 YYYYMMDD_HHMM
        current_time = datetime.now().strftime("%Y%m%d_%H%M")
        # 设置文档文件名为当前时间
        self.docfilename = f"output_{current_time}.docx"
        self.projectroot = os.path.dirname(os.path.abspath(__file__))
        self.docfilepath = os.path.join(
            self.projectroot, self.docfilename
        )
        self.doc = Document()
        
        """
        # 获取当前时间并将其转换为字符串
        current_time = datetime.now().strftime('%Y%m%d_%H%M%S')
        # 根据当前时间创建.log文件的名称
        self.logfilename = f"log_{current_time}.txt"
        # 获取当前目录的路径（也就是你的脚本所在的目录）
        self.logfileroot = os.path.dirname(os.path.abspath(__file__))
        # 将目录和文件名结合，获得完整的文件路径
        self.logfilename = os.path.join(
            self.logfileroot, self.logfilename
        )
        # 配置日志记录器
        logging.basicConfig(
            filename=self.logfilename,
            level=logging.DEBUG,
            format='%(asctime)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        # 使用日志记录器
        logging.info('This is an info message')
        logging.error('This is an error message')
        """

    def static_value(self):
        self.savelog = True
        self.seller = True
        self.country = None
        self.info_all_count = 0
        self.info_fail_count = 0
        self.asin_all_count = 0
        self.asin_fail_count = 0
        self.div_class_list = list()
        self.section_dict = {}

        self.sql = MysqlUtil()

        # 品牌广告
        # _c2Itd_cardContainer_27VO-
        self.result_class_SB = (
            's-result-item s-widget s-widget-spacing-large AdHolder s-flex-full-width'
        )
        # 自然结果
        #                       sg-col-4-of-24 sg-col-4-of-12 s-result-item s-asin sg-col-4-of-16 sg-col s-widget-spacing-small sg-col-4-of-20
        self.result_class_NR = 'sg-col-4-of-24 sg-col-4-of-12 s-result-item s-asin sg-col-4-of-16 sg-col s-widget-spacing-small sg-col-4-of-20'
        # 商品广告
        #                       sg-col-4-of-24 sg-col-4-of-12 s-result-item s-asin sg-col-4-of-16 AdHolder sg-col s-widget-spacing-small sg-col-4-of-20
        self.result_class_SP = 'sg-col-4-of-24 sg-col-4-of-12 s-result-item s-asin sg-col-4-of-16 AdHolder sg-col s-widget-spacing-small sg-col-4-of-20'
        # 高评价推荐 Trending now
        #                       sg-col-20-of-24 s-result-item sg-col-0-of-12 sg-col-16-of-20 s-widget sg-col s-flex-geom sg-col-12-of-16 s-widget-spacing-large
        self.result_class_HR = 'sg-col-20-of-24 s-result-item sg-col-0-of-12 sg-col-16-of-20 s-widget sg-col s-flex-geom sg-col-12-of-16 s-widget-spacing-large'
        # 视频广告
        #                       sg-col-20-of-24 s-result-item sg-col-0-of-12 sg-col-16-of-20 s-widget sg-col s-flex-geom s-widget-spacing-small sg-col-12-of-16
        self.result_class_BV = 'sg-col-20-of-24 s-result-item sg-col-0-of-12 sg-col-16-of-20 s-widget sg-col s-flex-geom s-widget-spacing-small sg-col-12-of-16'
        # Results
        self.result_class_Head = 'a-section a-spacing-none s-result-item s-flex-full-width s-border-bottom-none s-widget s-widget-spacing-large'
        self.result_class_Bottom = 'a-section a-spacing-none s-result-item s-flex-full-width s-widget s-widget-spacing-large'

    # 调试时，每次get_Search都运行一次
    def info_init(self):
        self.asin = None
        self.asin_url = None  # 点击后跳转的url链接？
        self.image = None
        self.title = None
        self.amz_choice = None
        self.amz_choice_type = None
        self.tag = None
        self.best_seller = None
        self.variant = None
        self.variant_type = None
        self.variant_count = None
        self.deal_type = None
        self.is_deal = None  # boolean value
        self.fba = None
        self.is_fba = None  # boolean value
        self.is_amz = None  # boolean value
        self.is_smb = None  # boolean value
        self.left_count = None  # TINYINT
        self.lower_price = None
        self.lower_list = None
        self.lowest_precent = None
        self.sale_price = None
        self.rrp_price = None
        self.rrp_type = None
        self.prime_price = None
        self.subscribe_price = None
        self.fba_price = None
        self.profit_rate = None
        self.discount = None
        self.coupon = None
        self.saving = None
        self.promotion = None
        self.promotion_code = None
        self.rating = None
        self.review = None
        self.bought = None
        self.a_rank_name = None
        self.a_rank = None
        self.b_rank_name = None
        self.b_rank = None
        self.c_rank_name = None
        self.c_rank = None
        self.d_rank_name = None
        self.d_rank = None

        self.shopname = None
        self.brand = None
        self.follow = None
        self.merchant_token = None
        self.use_age = None
        self.use_ages_do = None
        self.use_ages_up = None
        self.bullet_points_1 = None
        self.bullet_points_2 = None
        self.bullet_points_3 = None
        self.bullet_points_4 = None
        self.bullet_points_5 = None
        self.bullet_points_6 = None
        self.bullet_points = None
        self.base_info = None
        self.weight = None
        self.weight_unit = None
        self.length_l = None
        self.length_w = None
        self.length_h = None
        self.length_unit = None
        self.start_sale_time = None
        
        # 西柚数据
        self.xiyou_seven_days_views =None
        self.xiyou_na_ratio = None
        self.xiyou_ad_ratio = None

        # 实时变化
        self.data_index = None
        self.data_uuid = None
        self.data_component_type = None
        self.data_component_id = None
        self.data_cel_widget = None
        self.data_type = None

        self.time = None

    
    # 输出所有字典信息
    def to_all_dict(self):
        return {
            "asin": self.asin,
            "asin_url": self.asin_url,
            "image": self.image,
            "title": self.title,
            "amz_choice": self.amz_choice,
            "amz_choice_type": self.amz_choice_type,
            "tag": self.tag,
            "best_seller": self.best_seller,
            "variant": self.variant,
            "variant_type": self.variant_type,
            "variant_count": self.variant_count,
            "deal_type": self.deal_type,
            "is_deal": self.is_deal,
            "fba": self.fba,
            "is_fba": self.is_fba,
            "is_amz": self.is_amz,
            "is_smb": self.is_smb,
            "left_count": self.left_count,
            "lower_price": self.lower_price,
            "lower_list": self.lower_list,
            "lowest_precent": self.lowest_precent,
            "sale_price": self.sale_price,
            "rrp_price": self.rrp_price,
            "rrp_type": self.rrp_type,
            "prime_price": self.prime_price,
            "fba_price": self.fba_price,
            "profit_rate": self.profit_rate,
            "subscribe_price": self.subscribe_price,
            "discount": self.discount,
            "coupon": self.coupon,
            "saving": self.saving,
            "promotion": self.promotion,
            "rating": self.rating,
            "review": self.review,
            "bought": self.bought,
            "a_rank_name": self.a_rank_name,
            "a_rank": self.a_rank,
            "b_rank_name": self.b_rank_name,
            "b_rank": self.b_rank,
            "c_rank_name": self.c_rank_name,
            "c_rank": self.c_rank,
            "d_rank_name": self.d_rank_name,
            "d_rank": self.d_rank,
            "shopname": self.shopname,
            "brand": self.brand,
            "follow": self.follow,
            "merchant_token": self.merchant_token,
            "use_age": self.use_age,
            "use_ages_do": self.use_ages_do,
            "use_ages_up": self.use_ages_up,
            "bullet_points_1": self.bullet_points_1,
            "bullet_points_2": self.bullet_points_2,
            "bullet_points_3": self.bullet_points_3,
            "bullet_points_4": self.bullet_points_4,
            "bullet_points_5": self.bullet_points_5,
            "bullet_points_6": self.bullet_points_6,
            "bullet_points": self.bullet_points,
            "base_info": self.base_info,
            "weight": self.weight,
            "weight_unit": self.weight_unit,
            "length_l": self.length_l,
            "length_w": self.length_w,
            "length_h": self.length_h,
            "length_unit": self.length_unit,
            "start_sale_time": self.start_sale_time,
            "data_index": self.data_index,
            "data_uuid": self.data_uuid,
            "data_component_type": self.data_component_type,
            "data_component_id": self.data_component_id,
            "data_cel_widget": self.data_cel_widget,
            "data_type": self.data_type,
            "xiyou_seven_days_views":self.xiyou_seven_days_views,
            "xiyou_na_ratio":self.xiyou_na_ratio,
            "xiyou_ad_ratio":self.xiyou_ad_ratio,
            "time": self.time
        }

    # 输出数据库info_current的对应字典
    def to_info_current_dict(self):
        return {
            "asin": self.asin,
            "country": self.country,
            "time": self.time,
            "image": self.image,
            "title": self.title,
            "brand": self.brand,
            "shopname": self.shopname,
            "merchant_token": self.merchant_token,
            "follow": self.follow,
            "amz_choice": self.amz_choice,
            "amz_choice_type": self.amz_choice_type,
            "tag": self.tag,
            "best_seller": self.best_seller,
            "deal_type": self.deal_type,
            "is_deal": self.is_deal,
            "is_fba": self.is_fba,
            "is_amz": self.is_amz,
            "is_smb": self.is_smb,
            "left_count": self.left_count,
            "lower_price": self.lower_price,
            "lower_list": self.lower_list,
            "lowest_precent": self.lowest_precent,
            "sale_price": self.sale_price,
            "rrp_price": self.rrp_price,
            "rrp_type": self.rrp_type,
            "prime_price": self.prime_price,
            "fba_price": self.fba_price,
            "profit_rate": self.profit_rate,
            "subscribe_price": self.subscribe_price,
            "discount": self.discount,
            "coupon": self.coupon,
            "saving": self.saving,
            "promotion": self.promotion,
            "promotion_code": self.promotion_code,
            "rating": self.rating,
            "review": self.review,
            "bought": self.bought,
            "a_rank_name": self.a_rank_name,
            "a_rank": self.a_rank,
            "b_rank_name": self.b_rank_name,
            "b_rank": self.b_rank,
            "c_rank_name": self.c_rank_name,
            "c_rank": self.c_rank,
            "d_rank_name": self.d_rank_name,
            "d_rank": self.d_rank,
            "xiyou_seven_days_views":self.xiyou_seven_days_views,
            "xiyou_na_ratio":self.xiyou_na_ratio,
            "xiyou_ad_ratio":self.xiyou_ad_ratio,
        }

    def append_line(self, line, flag=True):
        # 当 savelog=True 且 flag=True 时，整个表达式为 True。
        # 当 savelog=True 且 flag=False 时，整个表达式为 False。
        # 当 savelog=False 且 flag=True 时，整个表达式为 True。
        # 当 savelog=False 且 flag=False 时，整个表达式为 False。
        # if self.savelog and flag or (not self.savelog and flag):
        if flag:
            # 追加一行到文档中
            self.doc.add_paragraph(line)
            # 保存文档
            self.doc.save(self.docfilepath)
        print(line)
        
    def append_dict(self, dict):        
        # 迭代字典并添加内容
        for key, value in dict.items():
            self.doc.add_paragraph(f"{key}: {value}")
        # 保存文档
        self.doc.save(self.docfilepath)

    # 搜索栏输入关键词
    def start_search(self, search_words):
        search_box = self.driver.find_element(
            By.XPATH, '//*[@id="twotabsearchtextbox"]'
        )
        search_box.send_keys(search_words)
        self.driver.find_element(By.ID, "nav-search-submit-button").click()
        self.append_line(f'输入关键词{search_words}')

    def roll_down(self):
        # 模拟滚轮操作，移动到页面底部
        last_elem = self.driver.find_element(By.XPATH,'//*[@id="rhf"]')
        # 模拟滚轮操作，移动到页面底部
        while True:
            is_display = last_elem.is_displayed()
            driver.execute_script("window.scrollBy(0, 100)")  # Scroll down 100px
            time.sleep(0.3)  # Waiting 0.3 sec
            if not is_display:
                break
        while True:
            origin_position = last_elem.location['y']  # Get the original position of the element
            driver.execute_script("window.scrollBy(0, 100)")  # Scroll down 100px
            time.sleep(0.3)  # Waiting 0.3 sec
            if origin_position == last_elem.location['y']:
                break  # If the position of the element has not changed, the page is at the bottom

    def garb_search(self):
        current_url = self.driver.current_url
        parsed_url = urlparse(current_url)
        lower_case_netloc = parsed_url.netloc.lower()
        if 'com' in lower_case_netloc:
            self.country = 'us'
        elif 'co.uk' in lower_case_netloc:
            self.country = 'uk'
        else:
            self.country = lower_case_netloc.split('.')[-1]
        
        self.wait.until(EC.presence_of_all_elements_located((By.XPATH, '//*[@id="search"]/div[contains(@class,"s-desktop-content")]/div[1]/div[2]/span[1]/div[1]')))
        div_list = self.driver.find_element(By.XPATH, ".//div[contains(@class, 's-result-list s-search-results')]").find_elements(By.XPATH, "./div")
        self.append_line(f'找到{len(div_list)}个搜索结果')

        for div in div_list:
            # 定位div的位置，方便后期调试使用
            self.append_line(self.get_xpath(div))
            # 初始化单个ASIN数据单元
            self.info_init()
            # 存储所有div的class
            div_class = div.get_attribute('class')
            self.div_class_list.append(div_class)
            # 获取基本属性信息
            self.get_index(div)

            # 品牌广告
            if div_class == self.result_class_SB:
                self.append_line(f'==品牌广告==\t{self.data_index}=\t{self.data_cel_widget}')
                self.asin_all_count += 1
                self.data_type = "SB"
                
            # 结果抬头
            elif div_class == self.result_class_Head:
                self.append_line(f'==结果抬头==\t{self.data_index}=\t{self.data_cel_widget}')
                
            # 自然位置
            elif div_class == self.result_class_NR:
                self.append_line(f'==自然位置==\t{self.data_index}=\t{self.data_cel_widget}')
                self.asin_all_count += 1
                self.data_type = "NR"
                self.get_search_data(div)
                self.format_data()
                # 使用default=str将所有非字符串对象转换为字符串
                # 显示抓取的数据
                # self.append_line(json.dumps(self.to_all_dict(), default=str))
                self.append_dict(self.to_all_dict())
                self.time = datetime.now().strftime("%Y-%m-%d")
                # 插入数据
                result = self.sql.replace_by_dict(
                    "asin_info_current", self.to_info_current_dict()
                )
                self.append_line(self.to_info_current_dict())
                self.append_line(f'##插入数据asin_info_current:{result}')
                self.append_line('##################################')
                
            # 广告位置
            elif div_class == self.result_class_SP:
                self.append_line(f'==广告位置==\t{self.data_index}=\t{self.data_cel_widget}')
                self.asin_all_count += 1
                self.data_type = "SP"
                self.get_search_data(div)
                self.format_data()
                # 显示抓取的数据
                # self.append_line(json.dumps(self.to_all_dict(), default=str))
                self.append_dict(self.to_all_dict())
                # 插入数据
                self.time = datetime.now().strftime("%Y-%m-%d")
                result = self.sql.replace_by_dict(
                    "asin_info_current", self.to_info_current_dict()
                )
                self.append_line(f'##插入数据asin_info_current:{result}')
                self.append_line('##################################')
                
            # 评价推荐
            elif div_class == self.result_class_HR:
                self.append_line(f'==评价推荐==\t{self.data_index}=\t{self.data_cel_widget}')
                lis = div.find_elements(By.TAG_NAME, 'li')
                inner_divs = []
                # 获取每个li中的div
                for li in lis:
                    inner_divs.append(li.find_element(By.XPATH, './div'))
                # 遍历上面获取到的div
                for div in inner_divs:
                    # 初始化
                    self.info_init
                    self.get_index(div)
                    self.asin_all_count += 1
                    self.data_type = "HR"
                    self.get_search_data(div)
                    self.format_data()
                    # 显示抓取的数据
                    # self.append_line(json.dumps(self.to_all_dict(), default=str))
                    self.append_dict(self.to_all_dict())
                    self.time = datetime.now().strftime("%Y-%m-%d")
                    # 插入数据
                    result = self.sql.replace_by_dict(
                        "asin_info_current", self.to_info_current_dict()
                    )
                    self.append_line(f'##插入数据asin_info_current:{result}')
                    self.append_line('##################################')
                    
            # 视频广告
            elif div_class == self.result_class_BV:
                self.append_line(f'==视频广告==\t{self.data_index}=\t{self.data_cel_widget}')
                self.asin_all_count += 1
                self.data_type = "BV"
                
            elif div_class == self.result_class_Bottom:
                self.append_line(f'==底部内容==\t{self.data_index}=\t{self.data_cel_widget}')
                
            else:
                self.append_line(f'=!=未知广告==\t{self.data_index}=\t{div_class}')

    # 抓取结束时统计抓取数据
    def stat_seart_result(self):
        # 去除重复元素
        self.div_class_list = DataType.remove_duplicates(self.div_class_list)
        self.append_line("================")
        self.append_line(self.div_class_list)
        self.append_line(
            f'info_all_count: {self.info_all_count}=>info_fail_count: {self.info_fail_count}'
        )
        self.append_line(
            f'asin_all_count: {self.asin_all_count}=>asin_fail_count: {self.asin_fail_count}'
        )

    # 获取div的特征值 
    def get_div_features(self, parent, level=0,flag=False):
        # 获取父元素的标签名
        tag_name = parent.tag_name
        # 获取父元素的class属性
        attributes = parent.get_attribute("class")
        # 获取父元素的父元素
        pre_parent = parent.find_element(By.XPATH, '..')
        # 获取父元素的所有同类型的同级元素
        elements = pre_parent.find_elements(By.TAG_NAME, parent.tag_name)
        i = 0
        # 递归时，当前步骤中的parent不是第一次输入的值时
        if level != 0:
            # 遍历同级元素列表
            for i, elem in enumerate(elements):
                # 当遍历到当前元素时，记录当前元素在同级元素中的位置
                if elem != parent:
                    i = i + 1
                # 当循环到自己时，停止循环，避免过度循环
                elif elem == parent:
                    break
        # 标签类型，层级，该层级下有多少个相关元素，class的值，[]下一级元素内容
        feature = [tag_name, level, i, attributes, []]
        if flag:
            space_interval = level*"  "
            self.append_line(f'{space_interval}{feature}')
        # 遍历当前div元素的子元素
        for child in parent.find_elements(By.XPATH, './*'):
            # 对每个子元素递归调用get_div_features()函数，获取子元素的特征
            child_feature = self.get_div_features(child, level + 1, flag)
            # 将子元素的特征添加到当前div元素的特征列表中
            feature[4].append(child_feature)
        return feature

    def get_element_structure(self, element, level=0):
        outer_html = element.get_attribute('outerHTML')
        soup = BeautifulSoup(outer_html, "html.parser")
        # 使用prettify()函数将HTML代码美化
        pretty_html = soup.prettify()
        # 打印美化后的HTML代码，每行打印一次
        for line in pretty_html.split('\n'):
            self.append_line(line)


    def cop_div_features(self, div):
        return

    # 比对div的特征值，找到配置文件中对应的section，在section中根据数据位置获取数据
    # 如果特征值在配置文件中没有找到，则保存特征值
    def match_feature_data(self, div_Info_child, config_name):
        config_file = 'yaml/features_result_'+config_name+'.yml'
        # 获取项目根目录路径
        root_folder = os.path.dirname(os.path.abspath(__file__))
        config_img = os.path.join(root_folder, 'yaml', f'img_{config_name}')
        
        self.append_line(f'&&当前ASIN下的div元素有：{len(div_Info_child)}个')
        with open(config_file) as f:
            config = yaml.safe_load(f)
        
        # 遍历div_Info_child数组
        for index, child in enumerate(div_Info_child, 1):
            child_class = child.get_attribute("class")
            feature_symbol = None
            if child_class == '':
                feature_symbol = 'div'
            child_xpath = self.get_xpath(child)
            self.append_line(f'--当前匹配的是第 {index} 个div')
            self.append_line(f'--当前匹配的xpath：{child_xpath}')
            # 先获取div的特征值
            features_list = self.get_div_features(child, 0)
            # 将特征值转为json
            feature_str = json.dumps(features_list)
            exists = False
            self.info_all_count += 1
            key_value = None
            value = None
            # 遍历配置文件
            for section in config:
                # 匹配特征值
                if config[section]['Div_feature'] == feature_str:
                    self.count_section(section)
                    self.append_line(f'--已匹配对应的特征值section：{section}')
                    exists = True
                    key_flag = 1
                    while key_flag > 0:
                        # section中有几个需要获取的data，后缀就到几
                        key_name = f'data_{key_flag}'
                        # 匹配失败，即section的需要获取值已经到尾部，跳出循环
                        if key_name not in config[section]:
                            key_flag = -1
                            continue
                        # 匹配成功后，获取键值
                        key_value = config[section][key_name]
                        data_name = key_value[0]
                        data_method = key_value[1]
                        data_xpath = key_value[2]
                        data_type = key_value[3]
                        target_div = child.find_element(By.XPATH, data_xpath)
                        # 获取值的方式
                        # 如果data_type=hiddentext，通过js代码修改元素的class，再获取值
                        if data_method == 'xpath':
                            if data_type == 'hiddentext':
                                self.driver.execute_script(
                                    "arguments[0].className = '';", target_div
                                )
                            value = target_div.text
                        # 获取元素的属性
                        if data_method == 'attribute':
                            value = target_div.get_attribute(data_type)
                        self.append_line(f'匹配到的：{section}\t_数据名称：{data_name}\t_数据值=>{value}')
                        # 通过data_name找到self中对应名称的值，将value值赋值给self.'data_name'
                        # 如果 self 之前没有叫做 data_name 的属性，这条语句将会给 self 增加一个新的属性，并将 value 赋值给它
                        setattr(self, data_name, value)
                        key_flag = key_flag + 1
                    break  # exists = True
            new_section = None
            if not exists:
                self.get_element_structure(child)
                # 重新获取一次div特征值，并打印在文档中
                self.get_div_features(child, 0, True)
                self.info_fail_count += 1
                # 'Section{}'.format(len(config)+1) len(config)+1 就是数量+1,即新增一个section后的总数，'Section{}'.format(3) = 'Section3'
                #new_section = f'{feature_symbol} data_i.{info.data_index} c.{index}'
                new_section = f'data_i.{self.data_index}_r.{self.data_cel_widget}_c.{index}'                # 当 data_index 为 None时，该div时HR下的asin
                config[new_section] = {'Div_feature': feature_str}
                self.append_line(f'!!新增元素特征值：{new_section}\t目标配置类型：{config_name}')
                with open(config_file, 'w') as f:
                    yaml.dump(config, f)
                    
                y = child.location['y']
                self.driver.execute_script(f"window.scrollTo(0, {y});")
                # 获取元素的x,y坐标
                child_x = child.location['x']
                child_y = child.location['y']
                child_width = child.size['width']
                child_height = child.size['height']
                # 截取指定区域
                self.driver.get_screenshot_as_file('screenshot.png')
                img = Image.open('screenshot.png')
                img = img.crop((child_x, 0, child_x+child_width, child_height))
                img.save(f'{config_img}\\{new_section}.png')

    # 获取基本属性信息
    def get_index(self, div):
        self.asin = div.get_attribute("data-asin")
        self.data_index = div.get_attribute("data-index")
        self.data_uuid = div.get_attribute("data-uuid")
        self.data_component_type = div.get_attribute("data-component-type")
        self.data_component_id = div.get_attribute("data-component-id")
        # value_if_true if condition else value_if_false
        self.data_cel_widget = (div.get_attribute('data-cel-widget').replace('search_result_', '') if div.get_attribute('data-cel-widget') is not None else '')
        if self.data_index == None:
            self.data_index = f'Z{div.find_element(By.XPATH, "..").get_attribute("aria-posinset")}'
        elif int(self.data_index) < 10:
            self.data_index = f'0{self.data_index}'

    # 获取div_info中的数据
    def get_search_data(self, div):
        #self.roll_down()
        
        # 找到 "a-section a-spacing-base" 为class的div
        # div_main = self.find_target_div_by_class(div, "a-section a-spacing-base")
        # 应该使用.//而不是//。因为//会在整个文档中查找，而.//则只在div1的子元素中查找。
        self.append_line(f'{self.get_xpath(div)}')
        #div_main = div.find_element(By.XPATH, ".//div[@class='a-section a-spacing-base']")
        div_main = div.find_element(By.XPATH, ".//div[contains(@class, 'a-section a-spacing-base')]")
        # a-section a-spacing-base a-text-center 在评价推荐中的ASINdiv中class可能包含a-text-center
        
        # 同级div下的class为none是卖家精灵 
        #div_main_sibling = div_main.find_elements(By.XPATH, "./following-sibling::div")
        div_main_sibling = div_main.find_element(By.XPATH, "..").find_elements(By.XPATH, "./div")
        #div_main_siblings = div_main.find_elements(By.XPATH, './following-sibling::div') + div_main.find_elements(By.XPATH, './preceding-sibling::div')
        div_seller = None
        div_xiyou = None
        for div in div_main_sibling:
            # print(div.get_attribute("class"))
            if div.get_attribute("class") == "xy_app":
                div_xiyou = div
            if div.get_attribute("class") == '':
                div_seller= div
        
        # 判断是否有Amazon'sChoice 或 Best Seller标志
        div_main_child = div_main.find_elements(By.XPATH, "./div")
        div_main_child_len = len(div_main_child)
        div_main_AB = None
        if div_main_child_len == 3:
            div_main_AB = div_main.find_element(By.XPATH, "./div[1]")
            div_main_Image = div_main.find_element(By.XPATH, "./div[2]")
            div_main_Info = div_main.find_element(By.XPATH, "./div[3]")
        if div_main_child_len == 2:
            div_main_Image = div_main.find_element(By.XPATH, "./div[1]")
            div_main_Info = div_main.find_element(By.XPATH, "./div[2]")
        # div_main_AB
        if div_main_AB:
            div_Info_AB_child = div_main_AB.find_elements(By.XPATH, "./*")
            self.match_feature_data(div_Info_AB_child, 'ab')
        # div_main_Image
        if div_main_Image:
            # 当ASIN所在的是评价推荐时，显示的结构会有部分不同
            try:
                asin_herf = div_main_Image.find_element(By.XPATH, "./span/a").get_attribute(
                    "href"
                )
                self.image = div_main_Image.find_element(
                    By.XPATH, "./span/a/div/img"
                ).get_attribute("src")
            except:
                asin_herf = div_main_Image.find_element(By.XPATH, "./div/span/a").get_attribute(
                    "href"
                )
                self.image = div_main_Image.find_element(
                    By.XPATH, "./div/span/a/div/img"
                ).get_attribute("src")
        # div_main_Info
        # 获取子元素，与features.yml的class特征进行配对，并抓取数据
        if div_main_Info:
            self.append_line(f'div_main_Info:{self.get_xpath(div_main_Info)}')
            #div_Info_child = div_main_Info.find_elements(By.XPATH, "./*")
            div_Info_child = div_main_Info.find_elements(By.XPATH, "./div")
            self.match_feature_data(div_Info_child, 'info')
            # 单独获取可售剩余数量
            self.get_left_count(div_main_Info, 'span', 'left in stock')
        # div_seller
        if div_seller:
            self.get_seller_data(div_seller)
        if div_xiyou:
            self.get_xiyou_data(div_xiyou)

    def get_seller_data(self, div_seller):
        self.append_line(f'{self.get_xpath(div_seller)}')
        # 获取Seller信息部分
        # 获取div_seller的outer HTML
        outer_html = div_seller.get_attribute("outerHTML")
        soup = BeautifulSoup(outer_html, "html.parser")
        # 查找包含特定文本的元素
        failure_message_element = soup.find("span", class_="loading-failed-tips")
        # 检查是否存在获取产品信息失败的文本
        if failure_message_element and "获取产品信息失败" in failure_message_element.text:
            self.append_line("@@不进行数据抓取，因为包含获取产品信息失败的文本")
            return
        
        # find 返回的是第一个匹配的元素，如果有多个满足条件的元素，它只返回第一个。空值返回 None
        # lambda 函数定义复杂的条件，确保 class 属性的值符合预期
        # div = soup.find("div", class_=lambda value: value and ("quick-view" in value) and ("quick-view-ext" in value))
        # 直接使用字符串匹配，精确某个特定的 class 属性值，不执行 lambda 函数的复杂条件判断，可能稍微快一些
        # div = soup.find("div", class_="quick-view quick-view-ext")
        # find_all 返回的是所有匹配的元素的列表，包括所有满足条件的元素。空值返回 []
        # div = soup.find_all("div", class_="quick-view quick-view-ext")
        # 品牌
        brand_span = soup.find('span', text=re.compile('品牌'))
        if brand_span:
            self.brand = brand_span.find_next_sibling('div').text
            self.append_line(f'卖家精灵：\t_数据名称：品牌\t_数据值=>{self.brand}')
        shopname_span = soup.find('span', class_=re.compile('word-title'), text=re.compile('卖家'))
        if shopname_span:
            self.shopname = shopname_span.find_next_sibling('span').text
            self.append_line(f'卖家精灵：\t_数据名称：店名\t_数据值=>{self.shopname}')
        delivery_span = soup.find('span', text=re.compile('配送'))
        if delivery_span:
            delivery_text = delivery_span.text.split(':')[0].strip(' ')
            if delivery_text == 'FBA':
                self.is_fba = 1
            elif delivery_text == 'FBM':
                self.is_fbm = 1
            elif delivery_text == 'AMZ':
                self.is_amz = 1
            self.append_line(f'卖家精灵：\t_数据名称：配送\t_数据值=>{delivery_text}')
        follow_span = soup.find('a', class_=re.compile('danger-tag'), text=re.compile('卖家: \d+'))
        if follow_span:
            self.follow = follow_span.text.split(':')[1].strip(' ')
            self.append_line(f'卖家精灵：\t_数据名称：卖家\t_数据值=>{self.follow}')
        # 排名
        rank_elements = soup.find_all("p", class_="bsr-list-item")
        for i, rank_elem in enumerate(rank_elements):
            category = rank_elem.find('span', class_='exts-color-blue').text
            number = rank_elem.find('span', class_='rank-box').text.strip('#')
            self.append_line(category, number)
            if i == 0:
                self.a_rank_name = category
                self.append_line(f'卖家精灵：\t_数据名称：排名1\t_数据值=>{self.a_rank_name}')
                self.a_rank = number
                self.append_line(f'卖家精灵：\t_数据名称：排名1值\t_数据值=>{self.a_rank}')
            elif i == 1:
                self.b_rank_name = category
                self.append_line(f'卖家精灵：\t_数据名称：排名1\t_数据值=>{self.b_rank_name}')
                self.b_rank = number
                self.append_line(f'卖家精灵：\t_数据名称：排名1值\t_数据值=>{self.b_rank}')
            elif i == 2:
                self.c_rank_name = category
                self.append_line(f'卖家精灵：\t_数据名称：排名1\t_数据值=>{self.c_rank_name}')
                self.c_rank = number
                self.append_line(f'卖家精灵：\t_数据名称：排名1值\t_数据值=>{self.c_rank}')
            elif i == 3:
                self.d_rank_name = category
                self.append_line(f'卖家精灵：\t_数据名称：排名1\t_数据值=>{self.d_rank_name}')
                self.d_rank = number
                self.append_line(f'卖家精灵：\t_数据名称：排名1值\t_数据值=>{self.d_rank}')
        # 价格
        fba_span = soup.find('span', string=re.compile('FBA费用'))
        if fba_span:
            self.fba_price = fba_span.find_next_sibling('span').text
            self.append_line(f'卖家精灵：\t_数据名称：FBA费用\t_数据值=>{self.fba_price}')
        profit_rate_span = soup.find('span', string=re.compile('毛利率'))
        if profit_rate_span:
            self.profit_rate = float(int(profit_rate_span.find_next_sibling('span').text.rstrip('%'))/100)
            self.append_line(f'卖家精灵：\t_数据名称：毛利率\t_数据值=>{self.profit_rate}')
        variant_count_span = soup.find('span', string=re.compile('变体数'))
        if variant_count_span:
            self.variant_count = variant_count_span.find_next_sibling('span').text
            self.append_line(f'卖家精灵：\t_数据名称：变体数\t_数据值=>{self.variant_count}')
        prime_price = soup.find('span', string=re.compile('Prime价格'))
        if prime_price:
            if self.prime_price == None:
                self.prime_price = prime_price.find_next_sibling('span').text
        # 重量
        weight_grams = soup.find('span', string=re.compile('grams'))
        weight_Kilograms = soup.find('span', string=re.compile('Kilograms'))
        weight_pounds = soup.find('span', string=re.compile('pounds'))
        weight_ounces = soup.find('span', string=re.compile('ounces'))
        if weight_grams:
            weight = weight_grams.text
            weight = re.search(r'(\d+\.?\d*)', weight).group(1)
            self.weight = int(weight)
            self.weight_unit = 'grams'
        elif weight_Kilograms:
            weight = weight_Kilograms.text
            weight = re.search(r'(\d+\.?\d*)', weight).group(1)
            self.weight = int(weight * 1000)
            self.weight_unit = 'grams'
        elif weight_pounds:
            weight = weight_pounds.text
            weight = re.search(r'(\d+\.?\d*)', weight).group(1)
            self.weight = int(weight * 1000 * 2.2046)
            self.weight_unit = 'grams'
        elif weight_ounces:
            weight = weight_ounces.text
            weight = re.search(r'(\d+\.?\d*)', weight).group(1)
            self.weight = int(weight * 1000 * 28.349)
            self.weight_unit = 'grams'
        self.append_line(f'卖家精灵：\t_数据名称：重量\t_数据值=>{self.weight}')
        self.append_line(f'卖家精灵：\t_数据名称：重量单位\t_数据值=>{self.weight_unit}')
        # 尺寸
        length_cm = soup.find('span', string=re.compile('cm'))
        length_inches = soup.find('span', string=re.compile('inches'))
        if length_cm:
            length = length_cm.text
            length_list = re.findall(r'(\d+\.?\d*)', length)
            length_list = [float(i) for i in length_list]
            length_list.sort(reverse=True)
            self.length_l = length_list[0]
            self.length_w = length_list[1]
            self.length_h = length_list[2]
            self.length_unit = 'cm'
        elif length_inches:
            length = length_inches.text
            length_list = re.findall(r'(\d+\.?\d*)', length)
            length_list = [float(i) for i in length_list]
            length_list.sort(reverse=True)
            self.length_l = length_list[0] * 2.54
            self.length_w = length_list[1] * 2.54
            self.length_h = length_list[2] * 2.54
            self.length_unit = 'cm'
        self.append_line(f'卖家精灵：\t_数据名称：长度\t_数据值=>{self.length_l}')
        self.append_line(f'卖家精灵：\t_数据名称：宽度\t_数据值=>{self.length_w}')
        self.append_line(f'卖家精灵：\t_数据名称：高度\t_数据值=>{self.length_h}')
        self.append_line(f'卖家精灵：\t_数据名称：长度单位\t_数据值=>{self.length_unit}')
        # 上架时间
        date_elme = soup.find('span', string=re.compile('上架时间'))
        if date_elme:
            date_span = date_elme.find_next_sibling("span")
            date_text = date_span.text
            date_match = re.search(r'(\d{4}-\d{2}-\d{2})', date_text)
            if date_match:
                date_str = date_match.group(1)
                self.start_sale_time = datetime.strptime(date_str, '%Y-%m-%d').date()
        self.append_line(f'卖家精灵：\t_数据名称：上架时间\t_数据值=>{self.start_sale_time}')
        # 五点描述
        # 详细信息
        
    def get_xiyou_data(self, div_xiyou):
        self.append_line(f'{self.get_xpath(div_xiyou)}')
        outer_html = div_xiyou.get_attribute("outerHTML")
        soup = BeautifulSoup(outer_html, "html.parser")
        # 提取数据
        xiyou_seven_days_views = soup.find('div', class_='num').text
        xiyou_na_ratio = soup.find('div', class_='left').text
        xiyou_ad_ratio = soup.find('div', class_='right').text
        # 格式数据
        self.xiyou_seven_days_views = int(xiyou_seven_days_views.replace(',', ''))
        self.xiyou_na_ratio = float(int(xiyou_na_ratio.rstrip('%'))/100)
        self.xiyou_ad_ratio = float(int(xiyou_ad_ratio.rstrip('%'))/100)
        self.append_line(f'西柚找词：\t_数据名称：七日流量\t_数据值=>{self.xiyou_seven_days_views}')
        self.append_line(f'西柚找词：\t_数据名称：自然流量\t_数据值=>{self.xiyou_na_ratio}')
        self.append_line(f'西柚找词：\t_数据名称：广告流量\t_数据值=>{self.xiyou_ad_ratio}')

    def format_data(self):
        if self.lower_price:
            self.lower_price = re.sub('[,$]', '', self.lower_price)
        if self.sale_price:
            self.sale_price = re.sub('[,$]', '', self.sale_price)
        if self.rrp_price:
            self.rrp_price = re.sub('[,$]', '', self.rrp_price)
        if self.prime_price:
            self.prime_price = re.sub('[,$]', '', self.prime_price)
        if self.fba_price:
            self.fba_price = re.sub('[,$]', '', self.fba_price)
        if self.a_rank:
            self.a_rank = self.a_rank.replace(",",".")
        if self.b_rank:
            self.b_rank = self.b_rank.replace(",",".")
        if self.c_rank:
            self.c_rank = self.c_rank.replace(",",".")
        if self.d_rank:
            self.d_rank = self.d_rank.replace(",",".")
        if self.rrp_type:
            self.rrp_type = self.rrp_type.strip(' ').strip(':')
        if self.start_sale_time:
            if isinstance(self.start_sale_time, str):
                try:
                    self.start_sale_time = datetime.strptime(self.start_sale_time, '%Y-%m-%d')
                except ValueError:
                    self.append_line(f"??时间字符串\"{self.start_sale_time}\"无法用\"%Y-%m-%d\"格式化")
                    self.start_sale_time = None
            elif hasattr(self.start_sale_time, 'strftime'):
                # Time is already in datetime format, no conversion needed
                pass
            else:
                self.append_line(f"??不清楚如何处理变量类型：{type(self.start_sale_time)}")
        if self.rating:
            self.rating = self.rating.split('out of')[0].strip(' ')
        if self.review:
            self.review = self.review.replace(",",".")
        if self.use_age:
            self.use_age = self.use_age.replace("Ages:",".").strip(' ')
        if self.bought:
            bought = self.bought.lower().split('+')[0]
            # 在字符串中第一次出现的位置（从0开始计数），如果字符串中不包含'k'，就返回-1。
            if bought.find('k')>0:
                self.bought = int(bought.split('k')[0]) * 1000
            elif bought.find('k')==-1:
                self.bought = int(bought.split('+')[0])
            else:
                self.bought = 0

    # 根据TAG_NAME及文本内容找到元素
    def get_tag_text_by_keyword(self, div, tag_name, keyword):
        spans = div.find_elements(By.TAG_NAME, tag_name)
        for span in spans:
            text = span.text
            if keyword in text:
                return text
        return None
    
    # 递归直到找到符合class_name的div
    def find_target_div_by_class(self, div, class_name):
        if div.get_attribute('class') == class_name:
            return div
        for child in div.find_elements(By.XPATH, "./div"):
            result = self.find_target_div_by_class(child, class_name)
            if result:
                return result
        return None

    def test_class_name(self, div_main_AB, div_main_Image, div_main_Info):
        # 测试三个div的class是否正确
        AB_class = "a-section a-spacing-none puis-status-badge-container aok-relative s-grid-status-badge-container puis-expand-height"
        Image_class = "s-product-image-container aok-relative s-text-center s-image-overlay-grey puis-image-overlay-grey"
        Info_class = (
            "a-section a-spacing-small puis-padding-left-small puis-padding-right-small"
        )
        if div_main_AB is not None:
            div_main_AB_class = div_main_AB.get_attribute("class")
            if div_main_AB_class == AB_class:
                self.append_line(f'{self.data_component_id},\tdiv_main_AB_class\t相等')
        div_main_Image_class = div_main_Image.get_attribute("class")
        if Image_class in div_main_Image_class:
            self.append_line(f'{self.data_component_id},\tdiv_main_Image_class\t包含')
        div_main_Info_class = div_main_Info.get_attribute("class")
        if div_main_Info_class == Info_class:
            self.append_line(f'{self.data_component_id},\tdiv_main_Info_class\t相等')

    # 单独获取剩余可售数量
    def get_left_count(self, div, tag_name, keyword):
        text = self.get_tag_text_by_keyword(div, 'span', 'left in stock')
        if text != None:
            self.left_count = re.findall(r'\d+', text)[0]
            # self.append_line(f'========{self.left_count}==========')

    # 计算section出现的次数
    def count_section(self, section_name):
        '''
        if section_name in self.section_dict:
            self.section_dict[section_name] += 1
        else:
            self.section_dict[section_name] = 1
        '''
        self.section_dict = defaultdict(int)
        self.section_dict[section_name] += 1

    def get_xpath(self, element):
    # 将JavaScript getXpath函数嵌入到Python代码中
        script = """
        function getElementXPath(element) {
            if (element && element.nodeType === Node.ELEMENT_NODE) {
                var xpath = '';
                var parent = null;
                for (parent = element; parent && parent !== document; parent = parent.parentNode) {
                    var index = 1;
                    for (previousSibling = parent.previousSibling; previousSibling; previousSibling = previousSibling.previousSibling) {
                        if (previousSibling.nodeName === parent.nodeName) {
                            index++;
                        }
                    }
                    xpath = '/' + parent.nodeName + '[' + index + ']' + xpath;
                }
                return xpath;
            } else {
                throw new Error('Invalid input: Element is not a valid DOM node');
            }
        }
        return getElementXPath(arguments[0]);
        """
        return self.driver.execute_script(script, element)
        # /HTML[1]/BODY[1]/DIV[3]/DIV[1]/DIV[2]/DIV[1]/DIV[2]/SPAN[1]/DIV[1]/DIV[3]


# 测试代码
if __name__ == '__main__':
    sc = ChromeStart("Seller", 9222)
    # sc.OpenPage("https://www.amazon.com/")
    sc.BindPage("https://www.amazon.com/s?k=", "Contain")
    driver, wait, actions = sc.GetDriver()
    AmazonS = AmazonSearch(driver, wait, actions)
    # AmazonS.garb_search("remote dinosaur")
    AmazonS.garb_search()

    # import cProfile
    # import pstats

    # cProfile.run("AmazonS.garb_search()", "stats.txt")
    # p = pstats.Stats('stats.txt')
    # p.strip_dirs().sort_stats(-1).self.append_line_stats()
